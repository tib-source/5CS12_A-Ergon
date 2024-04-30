import json
from django.contrib.auth.mixins import LoginRequiredMixin
from django.views.generic import TemplateView
from Bookings.models import Student, Equipment
from django.contrib.auth.models import User
from django.core.serializers.json import DjangoJSONEncoder
from django.shortcuts import render, redirect
from .forms import StudentRegisterForm
from django.http import JsonResponse
from django.utils import timezone
from datetime import timedelta
from django.http import FileResponse
from docx import Document
from reportlab.pdfgen import canvas
from io import BytesIO
import tempfile



class DashboardView(LoginRequiredMixin, TemplateView):
    template_name = 'Bookings/dashboard.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        # Retrieve the User object and add it to the context
        student = User.objects.get(pk=self.request.user.pk)  
        # Get all Equipment objects from the database
        equipment_queryset = Equipment.objects.all()

        # Convert queryset to a list of dictionaries
        serialized_data = list(equipment_queryset.values())

        # Convert choice field values to their human-readable representations
        for obj in serialized_data:
            obj['type'] = dict(Equipment.type_choices).get(obj['type'])
            obj['status'] = dict(Equipment.status_choices).get(obj['status'])

        # Serialize the data to JSON
        serialized_json = json.dumps(serialized_data, cls=DjangoJSONEncoder)

        # Add data into context object
        context['equipment'] = serialized_json
        context['student'] = student
        
        return context
    

# Registration View 
def student_register(response):
    if response.method == "POST":
        form = StudentRegisterForm(response.POST)
        if form.is_valid():
            form.save()
            return redirect("/home")
    else:
        form = StudentRegisterForm()

    return render(response, "registration/signup.html", {"form":form})

class ReportView(LoginRequiredMixin, TemplateView):
    template_name = 'report/report.html'

    def generate_report(self, report_type, time_period):
        """
        Generate report based on report type and time period.
        """
        # Calculate the start date based on the time period
        if time_period == 'last2weeks':
            start_date = timezone.now() - timedelta(days=14)
        elif time_period == 'last1month':
            start_date = timezone.now() - timedelta(days=30)
        elif time_period == 'last3months':
            start_date = timezone.now() - timedelta(days=90)
        elif time_period == 'last6months':
            start_date = timezone.now() - timedelta(days=180)
        elif time_period == 'last1year':
            start_date = timezone.now() - timedelta(days=365)
        elif time_period == 'last2years':
            start_date = timezone.now() - timedelta(days=730)
        else:
            # Default to last 1 month if time period is not recognized
            start_date = timezone.now() - timedelta(days=30)

        # Query equipment data based on the start date
        equipment_queryset = Equipment.objects.filter(last_audit__gte=start_date)

        # Convert queryset to a list of dictionaries
        serialized_data = list(equipment_queryset.values())

        # Convert choice field values to their readable representations
        for obj in serialized_data:
            obj['type'] = dict(Equipment.type_choices).get(obj['type'])
            obj['status'] = dict(Equipment.status_choices).get(obj['status'])

        #Generate report content based on report type
        if report_type == 'equipment':
            report_content = serialized_data
        else:
            report_content = "Report type not supported"

        return report_content

    def generate_pdf_report(self, report_data):
        """
        Generate PDF report.
        """
        buffer = BytesIO()
        p = canvas.Canvas(buffer)

        # Set initial y-coordinate for drawing text
        x = 100
        y = 750

        # Define maximum line width
        max_line_width = 75  

        # Define page height and bottom margin
        page_height = 800  
        bottom_margin = 50  

        # Draw the report data on the PDF canvas
        for entry in report_data:
            formatted_entry = ', '.join([f"{key}: {value if key != 'comment' or value else 'no comment'}" for key, value in entry.items()])
            # Split the formatted entry into segments that fit within the maximum line width
            segments = [formatted_entry[i:i+max_line_width] for i in range(0, len(formatted_entry), max_line_width)]

            # Calculate the remaining space on the current page
            remaining_space = y - bottom_margin

            # Check if there is enough space on the current page for the current entry
            if remaining_space < len(segments) * 15:
                # If not enough space, start a new page
                p.showPage()
                y = page_height - 50  # Reset y-coordinate for the new page

            # Draw each segment on the PDF canvas
            for segment in segments:
                p.drawString(100, y, segment)
                y -= 15

            # Move to the next line with an extra line spacing
            y -= 15 

        p.save()
        pdf_content = buffer.getvalue()
        buffer.close()

        return pdf_content

    def generate_doc_report(self, report_data):
        """
        Generate DOCX report.
        """
        doc = Document()

        # Draw the report data on the DOCX document
        for entry in report_data:
            formatted_entry = ', '.join([f"{key}: {value if key != 'comment' or value else 'no comment'}" for key, value in entry.items()])
            doc.add_paragraph(formatted_entry)

            # Add an empty paragraph for extra line spacing
            doc.add_paragraph()

        temp_file = tempfile.NamedTemporaryFile(delete=False)
        doc.save(temp_file)
        with open(temp_file.name, 'rb') as f:
            doc_content = f.read()

        return doc_content

    def post(self, request):
       # Retrieve form data
        report_type = request.POST.get('report_type')
        time_period = request.POST.get('time_period')
        email = request.POST.get('email')
        format_type = request.POST.get('format_type')

        try:
            # Generate report content
            report_data = self.generate_report(report_type, time_period)

            # Generate report based on format type
            if format_type == 'pdf':
                report_content = self.generate_pdf_report(report_data)
                content_type = 'application/pdf'
                file_extension = 'pdf'
            elif format_type == 'doc':
                report_content = self.generate_doc_report(report_data)
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                file_extension = 'docx'
            else:
                raise ValueError("Invalid format type. Must be 'pdf' or 'doc'.")

            # Prepare file response to download
            filename = f"report.{file_extension}"
            response = FileResponse(BytesIO(report_content), content_type=content_type)
            response['Content-Disposition'] = f'attachment; filename="{filename}"'

            return response
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)